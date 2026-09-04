from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PRO_USER_MANUAL.md"
OUTPUT = ROOT / "docs" / "Biashara_AI_Pro_User_Manual.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

FONT = "Calibri"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B677A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
GOLD = "9A6A00"
WHITE = "FFFFFF"
BLACK = "000000"
BORDER = "C8D2DF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color: str = BORDER, size: str = "8") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(col_widths_dxa)))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_dxa[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color: str = BORDER) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_num_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element

    abstract_ids = []
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        abstract_ids.append(int(abstract_num.get(qn("w:abstractNumId"))))
    next_abstract_id = max(abstract_ids or [0]) + 1

    num_ids = []
    for num in numbering.findall(qn("w:num")):
        num_ids.append(int(num.get(qn("w:numId"))))
    next_num_id = max(num_ids or [0]) + 1

    def make_abstract(abstract_id: int, fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")

        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        ppr.append(tabs)
        ppr.append(ind)

        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), FONT)
        rfonts.set(qn("w:hAnsi"), FONT)
        rpr.append(rfonts)

        for node in (start, num_fmt, lvl_text, lvl_jc, ppr, rpr):
            lvl.append(node)
        abstract.append(lvl)
        numbering.append(abstract)

    make_abstract(next_abstract_id, "bullet", "•")
    make_abstract(next_abstract_id + 1, "decimal", "%1.")

    def make_num(num_id: int, abstract_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    make_num(next_num_id, next_abstract_id)
    make_num(next_num_id + 1, next_abstract_id + 1)
    return next_num_id, next_abstract_id + 1


def add_num_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = []
    for num in numbering.findall(qn("w:num")):
        num_ids.append(int(num.get(qn("w:numId"))))
    next_num_id = max(num_ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return next_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = numpr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        numpr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    numid = numpr.find(qn("w:numId"))
    if numid is None:
        numid = OxmlElement("w:numId")
        numpr.append(numid)
    numid.set(qn("w:val"), str(num_id))


def configure_document(doc: Document) -> tuple[int, int]:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    bullet_num_id, decimal_abstract_id = add_num_definitions(doc)
    return bullet_num_id, decimal_abstract_id


def set_running_furniture(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    left = hp.add_run("Biashara AI Pro User Manual")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    hp.add_run("\t")
    right = hp.add_run("Pro Edition")
    set_run_font(right, size=8.5, color=MUTED)
    set_paragraph_border_bottom(hp, color="D7DBE2", size="6")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fp.add_run("Biashara AI Pro User Manual | "), size=8.5, color=MUTED)
    add_page_number(fp)
    for run in fp.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("OPERATOR MANUAL")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title.add_run("Biashara AI Pro\nUser Manual")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.line_spacing = 1.15
    run = subtitle.add_run("For business owners, managers, and staff using Biashara AI Pro")
    set_run_font(run, size=13.5, color=MUTED)

    rule = doc.add_paragraph()
    set_paragraph_border_bottom(rule, color=BLUE, size="18")
    rule.paragraph_format.space_after = Pt(18)

    summary = doc.add_paragraph()
    summary.paragraph_format.line_spacing = 1.25
    summary.paragraph_format.space_after = Pt(18)
    run = summary.add_run(
        "A practical guide to running product and service sales, stock, customers, credit, vouchers, ledger records, receipts, and on-device AI assistance from the phone."
    )
    set_run_font(run, size=11.5, color=BLACK)

    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="D7DBE2")
    widths = [2340, 2340, 2340, 2340]
    set_table_geometry(table, widths, indent_dxa=0)
    labels = ["Edition", "Audience", "Operating mode", "Documentation date"]
    values = ["Pro", "Owners, managers, staff", "Offline-first", "2026-05-20"]
    for idx, label in enumerate(labels):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_run_font(run, size=8.5, color=DARK_BLUE, bold=True)
    for idx, value in enumerate(values):
        set_cell_shading(table.cell(1, idx), WHITE)
        p = table.cell(1, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        set_run_font(run, size=9.5, color=BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    highlights = doc.add_table(rows=1, cols=1)
    highlights.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(highlights, [CONTENT_WIDTH_DXA], indent_dxa=0)
    set_table_borders(highlights, color="D7DBE2")
    cell = highlights.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("What Pro adds")
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(
        "Services catalogue, mixed product/service sales, staff assignment, prepaid vouchers, service QR verification, deposits, balance-due sales, and Pro service agents."
    )
    set_run_font(run, size=10.5, color=BLACK)

    doc.add_page_break()


def add_section_guide(doc: Document, markdown_lines: list[str]) -> None:
    h2s = [line[3:].strip() for line in markdown_lines if line.startswith("## ")]
    doc.add_heading("Section Guide", level=1)
    intro = doc.add_paragraph()
    intro.add_run("Use this guide as the table of contents for the manual. The sections follow the app workflow from setup through daily operation, reporting, troubleshooting, and reference.")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    set_table_geometry(table, [1800, 7560])
    headers = ["Section", "What it covers"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)

    for idx, title in enumerate(h2s, start=1):
        row = table.add_row()
        left = row.cells[0]
        right = row.cells[1]
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(lp.add_run(str(idx)), size=9.2, color=BLUE, bold=True)
        rp = right.paragraphs[0]
        set_run_font(rp.add_run(title), size=9.6, color=BLACK)
    doc.add_page_break()


def parse_inline(paragraph, text: str, base_size: float = 11, color: str = BLACK, bold_default: bool = False) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=color, bold=bold_default)


def clean_text(text: str) -> str:
    return text.strip().rstrip("  ")


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", line.strip()))


def table_widths(headers: list[str], cols: int) -> list[int]:
    joined = " ".join(headers).lower()
    if cols == 2:
        if "tab" in joined:
            return [1900, CONTENT_WIDTH_DXA - 1900]
        if "button" in joined:
            return [1750, CONTENT_WIDTH_DXA - 1750]
        if "agent" in joined:
            return [2300, CONTENT_WIDTH_DXA - 2300]
        if "term" in joined:
            return [2100, CONTENT_WIDTH_DXA - 2100]
        if "task" in joined:
            return [2700, CONTENT_WIDTH_DXA - 2700]
        return [2200, CONTENT_WIDTH_DXA - 2200]
    return [CONTENT_WIDTH_DXA // cols for _ in range(cols)]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    widths = table_widths(rows[0], cols)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    set_table_geometry(table, widths)

    for idx, text in enumerate(rows[0]):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and cols <= 2 else WD_ALIGN_PARAGRAPH.LEFT
        parse_inline(p, text, base_size=9.5, color=DARK_BLUE, bold_default=True)

    for row_values in rows[1:]:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            if idx == 0 and cols <= 2 and len(text) < 24:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parse_inline(p, text, base_size=9.4, color=BLACK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    parse_inline(p, text)


def add_numbered(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    parse_inline(p, text)


def is_note_text(text: str) -> bool:
    starts = (
        "Basic/Shop remains",
        "Use cost price carefully.",
        "Be careful.",
        "Credit requires a customer.",
        "The ledger is append-only.",
        "Always review OCR output.",
        "This is not the same as automatic",
        "Current limitation:",
        "Protect the phone",
        "Use adjustment entries",
    )
    return text.startswith(starts)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D7DBE2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    parse_inline(p, text, base_size=10.2, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_body_paragraph(doc: Document, text: str) -> None:
    text = clean_text(text)
    if not text:
        return
    if is_note_text(text):
        add_callout(doc, text)
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    if text.endswith(":") and len(text) <= 70:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        parse_inline(p, text, base_size=10.8, color=DARK_BLUE, bold_default=True)
    else:
        p.paragraph_format.space_after = Pt(6)
        parse_inline(p, text)


def parse_table_lines(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if is_table_separator(line):
            continue
        stripped = line.strip()
        if stripped.startswith("|"):
            stripped = stripped[1:]
        if stripped.endswith("|"):
            stripped = stripped[:-1]
        rows.append([clean_text(cell) for cell in stripped.split("|")])
    return rows


def add_manual_body(doc: Document, markdown: str, bullet_num_id: int, decimal_abstract_id: int) -> None:
    lines = markdown.splitlines()
    i = 0
    skip_front_matter = True
    while i < len(lines):
        line = lines[i].rstrip()

        if skip_front_matter:
            if line.startswith("## "):
                skip_front_matter = False
            else:
                i += 1
                continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_markdown_table(doc, parse_table_lines(table_lines))
            continue

        bullet = re.match(r"^\s*-\s+(.*)$", line)
        if bullet:
            while i < len(lines):
                match = re.match(r"^\s*-\s+(.*)$", lines[i].rstrip())
                if not match:
                    break
                add_bullet(doc, clean_text(match.group(1)), bullet_num_id)
                i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            decimal_num_id = add_num_instance(doc, decimal_abstract_id)
            while i < len(lines):
                match = re.match(r"^\s*\d+\.\s+(.*)$", lines[i].rstrip())
                if not match:
                    break
                add_numbered(doc, clean_text(match.group(1)), decimal_num_id)
                i += 1
            continue

        para_lines = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip() or nxt.startswith("## ") or nxt.startswith("### ") or nxt.startswith("|") or re.match(r"^\s*[-]\s+", nxt) or re.match(r"^\s*\d+\.\s+", nxt):
                break
            para_lines.append(nxt.strip())
            i += 1
        add_body_paragraph(doc, " ".join(para_lines))


def audit_tables(docx_path: Path) -> None:
    import zipfile
    from xml.etree import ElementTree as ET

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(docx_path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    for idx, tbl in enumerate(root.findall(".//w:tbl", ns), start=1):
        tblw = tbl.find("./w:tblPr/w:tblW", ns)
        grid = tbl.findall("./w:tblGrid/w:gridCol", ns)
        if tblw is None or not grid:
            raise RuntimeError(f"Table {idx} is missing explicit geometry")
        grid_sum = sum(int(col.attrib[f"{{{ns['w']}}}w"]) for col in grid)
        tbl_width = int(tblw.attrib[f"{{{ns['w']}}}w"])
        if grid_sum != tbl_width:
            raise RuntimeError(f"Table {idx} geometry mismatch: grid {grid_sum}, tblW {tbl_width}")


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    markdown_lines = markdown.splitlines()
    doc = Document()
    bullet_num_id, decimal_abstract_id = configure_document(doc)
    set_running_furniture(doc)
    add_cover(doc)
    add_section_guide(doc, markdown_lines)
    add_manual_body(doc, markdown, bullet_num_id, decimal_abstract_id)
    doc.core_properties.title = "Biashara AI Pro User Manual"
    doc.core_properties.subject = "Operator manual for Biashara AI Pro"
    doc.core_properties.keywords = "Biashara AI Pro, user manual, services, POS, ledger, vouchers, AI agents"
    doc.core_properties.comments = "Generated from docs/PRO_USER_MANUAL.md"
    doc.save(OUTPUT)
    audit_tables(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
